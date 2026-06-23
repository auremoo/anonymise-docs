#!/usr/bin/env python3
# Anonymize Docs — Auteur : Aurélien Moote - Moo - 2026 — Licence MIT
"""
Pipeline d'anonymisation hybride v2 : Regex + LLM local (Ollama /api/chat)
Usage CLI  : python anonymize.py <fichier> [--output fichier_sortie.md]
Usage Web  : streamlit run app.py

Dépendances :
  pip install requests python-docx pymupdf
"""

import io
import re
import sys
import json
import time
import argparse
import datetime
import threading
from pathlib import Path
from collections import defaultdict
from typing import Callable

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz  # PyMuPDF
    HAS_PDF = True
except ImportError:
    HAS_PDF = False


# =============================================================================
# LOGGER — Tout est tracé
# =============================================================================

class Logger:
    """Collecte tous les événements pour le rapport final."""

    def __init__(self, on_progress: Callable[[str, float], None] | None = None):
        self.entries = []
        self.start_time = time.time()
        self.on_progress = on_progress
        self.stats = {
            "fichier_source": "",
            "taille_originale": 0,
            "taille_finale": 0,
            "custom_remplacements": 0,
            "regex_remplacements": 0,
            "llm_passes": 0,
            "llm_chunks_traites": 0,
            "llm_erreurs": 0,
            "images_trouvees": 0,
            "verif_warnings": [],
            "duree_totale": 0,
            "annule": False,
        }

    def elapsed(self) -> str:
        return f"{time.time() - self.start_time:.0f}s"

    def log(self, level: str, message: str, progress: float | None = None):
        timestamp = time.time() - self.start_time
        entry = {"t": round(timestamp, 2), "level": level, "msg": message}
        self.entries.append(entry)
        # Affichage console
        icons = {"INFO": "📄", "REGEX": "🔍", "CUSTOM": "🏷️", "LLM": "🤖",
                 "VERIF": "🔎", "OK": "✅", "WARN": "⚠️", "ERROR": "❌",
                 "DONE": "✅", "STOP": "🛑", "IMG": "🖼️"}
        icon = icons.get(level, "•")
        print(f"  {icon} [{timestamp:6.1f}s] {message}")
        if self.on_progress and progress is not None:
            self.on_progress(f"{message} — {self.elapsed()}", progress)

    def generate_report(self, regex_anon, llm_entity_map,
                        images_folder: str = "") -> str:
        """Génère le rapport Markdown complet."""
        self.stats["duree_totale"] = round(time.time() - self.start_time, 2)
        now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        lines = [
            "# Rapport d'anonymisation",
            "",
            f"**Date :** {now}  ",
            f"**Fichier source :** `{self.stats['fichier_source']}`  ",
            f"**Taille originale :** {self.stats['taille_originale']} caractères  ",
            f"**Taille finale :** {self.stats['taille_finale']} caractères  ",
            f"**Durée totale :** {self.stats['duree_totale']}s  ",
        ]

        if self.stats["annule"]:
            lines.append("**Statut : ANNULÉ par l'utilisateur**  ")

        if self.stats["images_trouvees"] > 0:
            lines.append(
                f"**Images extraites :** {self.stats['images_trouvees']}  "
            )
            if images_folder:
                lines.append(f"**Dossier images :** `{images_folder}`  ")
            lines.append("")
            lines.append(
                f"> ⚠️ **{self.stats['images_trouvees']} image(s) extraite(s) "
                "et sauvegardée(s) dans un dossier séparé.** "
                "Les placeholders `[IMAGE_N]` dans le texte anonymisé "
                "correspondent aux fichiers numérotés dans le dossier. "
                "Vérifiez manuellement chaque image avant de la partager "
                "(logos, signatures, captures d'écran avec données visibles)."
            )

        lines += ["", "---", ""]

        # Table mots custom
        if self.stats["custom_remplacements"] > 0:
            lines.append("## Passe 0 — Mots personnalisés")
            lines.append("")
            lines.append(
                f"**{self.stats['custom_remplacements']} remplacement(s)**"
            )
            lines.append("")

        # Table regex
        lines.append("## Passe 1 — Regex (patterns structurés)")
        lines.append("")
        lines.append(f"**{self.stats['regex_remplacements']} remplacement(s)**")
        lines.append("")

        if regex_anon.mapping:
            lines.append("| Tag | Catégorie | Valeur originale |")
            lines.append("|-----|-----------|-----------------|")
            for key, tag in sorted(
                regex_anon.mapping.items(), key=lambda x: x[1]
            ):
                cat, original = key.split("::", 1)
                lines.append(f"| `{tag}` | {cat} | `{original}` |")
            lines.append("")

        # Table LLM
        lines.append("## Passe 2 & 3 — LLM (entités nommées)")
        lines.append("")
        lines.append(f"**Passes LLM :** {self.stats['llm_passes']}  ")
        lines.append(
            f"**Chunks traités :** {self.stats['llm_chunks_traites']}  "
        )
        lines.append(f"**Erreurs LLM :** {self.stats['llm_erreurs']}  ")
        lines.append("")

        if llm_entity_map:
            lines.append("**Entités détectées par le LLM :**")
            lines.append("")
            lines.append("| Tag | Occurrences |")
            lines.append("|-----|------------|")
            for tag, count in sorted(llm_entity_map.items()):
                lines.append(f"| `{tag}` | {count} |")
            lines.append("")

        # Vérification
        lines.append("## Vérification post-anonymisation")
        lines.append("")
        if self.stats["verif_warnings"]:
            for w in self.stats["verif_warnings"]:
                lines.append(f"- ⚠️ {w}")
        else:
            lines.append("✅ Aucun pattern résiduel détecté.")
        lines.append("")

        # Log complet
        lines += ["---", "", "## Journal complet", "", "```"]
        for e in self.entries:
            lines.append(f"[{e['t']:6.1f}s] [{e['level']:5s}] {e['msg']}")
        lines.append("```")

        return "\n".join(lines)


# =============================================================================
# PASSE 0 : Mots personnalisés
# =============================================================================

def apply_custom_words(
    text: str,
    custom_words: dict[str, str],
    regex_anon: "RegexAnonymizer",
) -> tuple[str, int]:
    """Remplace les mots custom avant la passe regex."""
    count = 0
    for word, category in sorted(custom_words.items(), key=lambda x: -len(x[0])):
        if not word.strip():
            continue
        tag = regex_anon._get_tag(category.upper(), word)
        text, n = re.subn(re.escape(word), tag, text, flags=re.IGNORECASE)
        count += n
    return text, count


# =============================================================================
# DICTIONNAIRE PERSISTANT
# =============================================================================

DICT_PATH = Path(__file__).resolve().parent / "sensitive-words.json"


def load_sensitive_words(filepath: Path | None = None) -> dict[str, str]:
    """Charge un dictionnaire de mots sensibles.
    Format fichier : {"CATEGORIE": ["mot1", "mot2"], ...}
    Retourne : {mot: catégorie} pour apply_custom_words().
    """
    path = filepath or DICT_PATH
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        words: dict[str, str] = {}
        for category, word_list in data.items():
            for word in word_list:
                if word.strip():
                    words[word.strip()] = category.upper()
        return words
    except Exception:
        return {}


def save_sensitive_words(
    words: dict[str, str], filepath: Path | None = None,
):
    """Sauvegarde le dictionnaire de mots sensibles.
    Entrée : {mot: catégorie}
    Format fichier : {"CATEGORIE": ["mot1", "mot2"], ...}
    """
    path = filepath or DICT_PATH
    grouped: dict[str, list[str]] = defaultdict(list)
    for word, category in sorted(words.items()):
        if word.strip():
            grouped[category.upper()].append(word.strip())
    path.write_text(
        json.dumps(dict(grouped), indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


# =============================================================================
# PASSE 1 : Regex
# =============================================================================

class RegexAnonymizer:

    def __init__(self):
        self.counters = defaultdict(int)
        self.mapping = {}

    def _get_tag(self, category: str, original: str) -> str:
        key = f"{category}::{original.strip().lower()}"
        if key not in self.mapping:
            self.counters[category] += 1
            self.mapping[key] = f"[{category}_{self.counters[category]}]"
        return self.mapping[key]

    def anonymize(self, text: str) -> str:
        # ── IP v4 ──
        text = re.sub(
            r'\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}'
            r'(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b',
            lambda m: self._get_tag("IP", m.group(0)), text)
        # ── IP v6 ──
        text = re.sub(
            r'\b(?:[0-9a-fA-F]{1,4}:){2,7}[0-9a-fA-F]{1,4}\b',
            lambda m: self._get_tag("IP", m.group(0)), text)
        # ── FQDN ──
        text = re.sub(
            r'\b[a-zA-Z][a-zA-Z0-9\-]*\.'
            r'(?:[a-zA-Z0-9\-]+\.)*'
            r'(?:local|lan|internal|corp|intra|net|com|fr|org|eu|io|'
            r'de|uk|it|es)\b',
            lambda m: self._get_tag("SERVEUR", m.group(0)), text)
        # ── Email ──
        text = re.sub(
            r'\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b',
            lambda m: self._get_tag("EMAIL", m.group(0)), text)
        # ── Credentials / connection strings ──
        # Password=xxx, Pwd=xxx, User ID=xxx dans les connection strings
        text = re.sub(
            r'(?i)((?:Password|Pwd|User\s*ID|Uid)\s*=\s*)([^;"\r\n]+)',
            lambda m: m.group(1) + self._get_tag("SECRET", m.group(2)),
            text)
        # Champs JSON sensibles : "ApiKey": "xxx"
        text = re.sub(
            r'(?i)("(?:Password|Pwd|Secret|ApiKey|api_key|Token|'
            r'AccessToken|ClientSecret|client_secret|AccessKey|'
            r'SecretKey|PrivateKey|bindCredentials)'
            r'"\s*:\s*")([^"]+)"',
            lambda m: m.group(1) + self._get_tag("SECRET", m.group(2))
            + '"',
            text)
        # ── Dates (AVANT téléphone pour éviter 29-01-2026 → TEL) ──
        text = re.sub(
            r'\b\d{1,2}[/.\-]\d{1,2}[/.\-]\d{2,4}\b',
            lambda m: self._get_tag("DATE", m.group(0)), text)
        text = re.sub(
            r'\b\d{4}-\d{2}-\d{2}\b',
            lambda m: self._get_tag("DATE", m.group(0)), text)
        # French written dates: "4 février 2026", "1er mars 2025"
        text = re.sub(
            r'\b\d{1,2}(?:er)?\s+(?:janvier|février|mars|avril|mai|juin|'
            r'juillet|août|septembre|octobre|novembre|décembre)\s+\d{4}\b',
            lambda m: self._get_tag("DATE", m.group(0)), text,
            flags=re.IGNORECASE)
        # English written dates: "4 February 2026", "March 1, 2025"
        text = re.sub(
            r'\b\d{1,2}\s+(?:January|February|March|April|May|June|'
            r'July|August|September|October|November|December)\s+\d{4}\b',
            lambda m: self._get_tag("DATE", m.group(0)), text,
            flags=re.IGNORECASE)
        text = re.sub(
            r'\b(?:January|February|March|April|May|June|'
            r'July|August|September|October|November|December)'
            r'\s+\d{1,2},?\s+\d{4}\b',
            lambda m: self._get_tag("DATE", m.group(0)), text,
            flags=re.IGNORECASE)
        # ── Téléphone (strict : doit commencer par 0X, +XX, ou 00XX) ──
        # Séparateurs = espace, point, tiret (PAS newline/tab)
        # Pattern 1 : FR  0X XX XX XX XX
        # Pattern 2 : Int +XX X XX XX XX XX
        # Pattern 3 : Int 00XX X XX XX XX XX (séparateur requis après indicatif)
        text = re.sub(
            r'\b0[1-9](?:[ .\-]?\d{2}){4}\b'
            r'|\+\d{1,3}[ .\-]?\(?\d{1,4}\)?(?:[ .\-]?\d{2,4}){2,5}\b'
            r'|\b00\d{2,3}[ .\-]\(?\d{1,4}\)?(?:[ .\-]?\d{2,4}){2,5}\b',
            lambda m: self._get_tag("TEL", m.group(0))
            if 8 <= len(re.sub(r'\D', '', m.group(0))) <= 15
            else m.group(0),
            text)
        # ── Chemins UNC ──
        text = re.sub(
            r'\\\\[a-zA-Z0-9\-_.]+(?:\\[a-zA-Z0-9\-_. ]+)+',
            lambda m: self._get_tag("CHEMIN", m.group(0)), text)
        # ── Chemins Linux /home ──
        text = re.sub(
            r'/home/[a-zA-Z][a-zA-Z0-9._\-]+',
            lambda m: self._get_tag("CHEMIN", m.group(0)), text)
        return text


# =============================================================================
# PASSE 2 + 3 : LLM via Ollama /api/chat
# =============================================================================

SYSTEM_PROMPT_PASS2 = """Tu es un outil d'anonymisation de documents techniques (industriel, IT/OT, cybersécurité). Remplace les entités nommées par des tags. Ne fais RIEN d'autre.

Lis d'abord le texte en entier pour repérer toutes les entités, puis remplace-les de manière cohérente.

CATÉGORIES À REMPLACER :
- Noms de personnes (prénoms, noms, initiales "J. Dupont") → [PERSONNE_1], [PERSONNE_2]...
- Entreprises, organisations, clients, prestataires → [ENTREPRISE_1], [ENTREPRISE_2]...
- Sites, usines, bâtiments, agences → [SITE_1], [SITE_2]...
- Projets internes, noms de code → [PROJET_1], [PROJET_2]...
- Adresses, villes, rues, codes postaux → [LIEU_1], [LIEU_2]...
- Numéros de contrat, références client, n° de commande → [REF_1], [REF_2]...

RÈGLES :
1. Même entité = même tag partout (ex: "Dupont" → toujours [PERSONNE_1]).
2. NE modifie RIEN d'autre : pas de reformulation, correction, ajout ou suppression.
3. Préserve le formatage : markdown, listes, tableaux, retours à la ligne.
4. Tags existants ([IP_1], [EMAIL_1], [DATE_1], [TEL_1], [SERVEUR_1], [CHEMIN_1], [SECRET_1], [IMAGE_1]...) → INTACTS.
5. En cas de doute → NE remplace PAS.

EXEMPLES :

Entrée : "Jean Dupont de la société Acme a visité l'usine de Lyon le [DATE_1]."
Sortie : "[PERSONNE_1] de la société [ENTREPRISE_1] a visité l'usine de [LIEU_1] le [DATE_1]."

Entrée : "Pour toute question, contactez Pierre à Strasbourg ou écrivez à [EMAIL_1]."
Sortie : "Pour toute question, contactez [PERSONNE_1] à [LIEU_1] ou écrivez à [EMAIL_1]."

Entrée : "J.D. de Sogetrel, contrat N°ABC-2024-0456, a réalisé l'audit du site de Fos-sur-Mer."
Sortie : "[PERSONNE_1] de [ENTREPRISE_1], contrat [REF_1], a réalisé l'audit du site de [LIEU_1]."

Entrée : "Cordialement,\\nMarie Lefevre\\nNexans — site de Bourg-en-Bresse\\nProjet INDUS-2025"
Sortie : "Cordialement,\\n[PERSONNE_1]\\n[ENTREPRISE_1] — site de [LIEU_1]\\n[PROJET_1]"

NE SONT PAS DES ENTITÉS (ne pas remplacer) :
- Marques technologiques utilisées comme termes techniques :
  Siemens, Schneider Electric, Rockwell, ABB, Honeywell, Yokogawa, Emerson, Beckhoff,
  Cisco, Microsoft, VMware, Fortinet, Palo Alto, Veeam, Acronis, Dell, HP, Lenovo
- Logiciels et protocoles :
  SCADA, WinCC, TIA Portal, Step 7, PCS 7, Unity Pro, FactoryTalk, RSLogix,
  OPC UA, OPC DA, Modbus, Profinet, Profibus, EtherNet/IP, BACnet, DNP3, IEC 61850,
  SQL Server, Windows Server, Active Directory, VMware ESXi, Hyper-V, Linux, RHEL,
  Python, Docker, Kubernetes, Ansible, Terraform, Git, Jenkins, Grafana, Zabbix, PRTG
- Termes métier OT/IT :
  PLC, RTU, HMI, DCS, SIS, MES, ERP, CMMS, GMAO, SNMP, VPN, DMZ, VLAN, firewall,
  automate, variateur, IHM, superviseur, historian, API REST, base de données

Retourne UNIQUEMENT le texte anonymisé. Aucun commentaire, aucune explication."""

SYSTEM_PROMPT_PASS3 = """Tu es un vérificateur d'anonymisation de documents techniques. Le texte a déjà été partiellement anonymisé mais il peut rester des oublis.

CHERCHE SPÉCIFIQUEMENT ces oublis fréquents :
- Prénoms isolés ("contactez Pierre", "signé Marie")
- Initiales ("J.D.", "M. Martin")
- Noms dans des adresses email partielles ou signatures
- Noms de villes ou rues non détectés ("basé à Strasbourg", "rue du Commerce")
- Références internes (n° contrat, n° client, n° badge, n° affaire)
- Noms de sociétés/clients non standard qui auraient été manqués

POUR CHAQUE OUBLI TROUVÉ :
- Utilise le tag suivant libre (ex: s'il y a déjà [PERSONNE_1] et [PERSONNE_2], utilise [PERSONNE_3])
- Si l'entité correspond à un tag existant, réutilise ce tag

EXEMPLES DE CORRECTIONS :

Avant : "[PERSONNE_1] a contacté Marc pour valider le projet."
Après : "[PERSONNE_1] a contacté [PERSONNE_2] pour valider le projet."

Avant : "Livraison prévue à Toulouse, bâtiment B3, pour le client Nexans."
Après : "Livraison prévue à [LIEU_1], bâtiment [SITE_1], pour le client [ENTREPRISE_1]."

NE TOUCHE PAS :
- Aux tags déjà en place ([PERSONNE_1], [IP_1], [LIEU_1], [ENTREPRISE_1], [REF_1], [SECRET_1], [IMAGE_1]... etc.)
- Aux termes techniques (SCADA, WinCC, OPC UA, PLC, Siemens, Schneider, etc.)
- Au formatage, à la structure, à la ponctuation

Si le texte est déjà bien anonymisé, retourne-le TEL QUEL, caractère pour caractère.

Retourne UNIQUEMENT le texte. Aucun commentaire."""


def call_ollama_chat(text: str, system_prompt: str, model: str = "gpt-oss:20b",
                     base_url: str = "http://localhost:11434",
                     timeout: int = 300) -> tuple[str, bool]:
    if not HAS_REQUESTS:
        return text, False
    try:
        response = requests.post(
            f"{base_url}/api/chat",
            json={
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user",
                     "content": f"Anonymise ce texte :\n\n{text}"}
                ],
                "stream": False,
                "options": {
                    "temperature": 0.05,
                    "top_p": 0.9,
                    "num_predict": 8192,
                    "num_ctx": 32768,
                }
            },
            timeout=timeout
        )
        response.raise_for_status()
        data = response.json()
        result = data.get("message", {}).get("content", "")
        if not result.strip():
            return text, False
        result = result.strip()
        # Strip markdown code blocks that some models wrap around output
        if result.startswith("```") and result.endswith("```"):
            lines = result.split("\n")
            if len(lines) >= 3:
                result = "\n".join(lines[1:-1]).strip()
        return result, True
    except (requests.exceptions.ConnectionError, requests.exceptions.Timeout):
        return text, False
    except Exception:
        return text, False


# =============================================================================
# LECTURE DE FICHIERS + EXTRACTION D'IMAGES
# =============================================================================

# Word XML namespaces
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def _extract_docx_image_rels(doc) -> dict[str, tuple[bytes, str]]:
    """Map relationship IDs to (image_bytes, extension) from a docx."""
    image_rels = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            blob = rel.target_part.blob
            ct = rel.target_part.content_type or "image/png"
            ext = ct.split("/")[-1]
            if ext == "jpeg":
                ext = "jpg"
            image_rels[rel_id] = (blob, ext)
    return image_rels


def _read_docx_with_images(data: bytes) -> tuple[str, list[tuple[bytes, str]]]:
    """Read docx → (text with [IMAGE_N] placeholders, [(bytes, ext), ...])."""
    doc = Document(io.BytesIO(data))
    image_rels = _extract_docx_image_rels(doc)
    images: list[tuple[bytes, str]] = []
    img_counter = 0
    parts: list[str] = []

    for para in doc.paragraphs:
        para_parts: list[str] = []
        for child in para._element:
            tag = child.tag
            # <w:r> — run element
            if tag == f"{{{_W_NS}}}r":
                has_image = False
                for blip in child.iter(f"{{{_A_NS}}}blip"):
                    embed = blip.get(f"{{{_R_NS}}}embed")
                    if embed and embed in image_rels:
                        img_counter += 1
                        images.append(image_rels[embed])
                        para_parts.append(f" [IMAGE_{img_counter}] ")
                        has_image = True
                if not has_image:
                    for t_elem in child.iter(f"{{{_W_NS}}}t"):
                        if t_elem.text:
                            para_parts.append(t_elem.text)
        parts.append("".join(para_parts))

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)

    return "\n".join(parts), images


def _read_pdf_with_images(data: bytes) -> tuple[str, list[tuple[bytes, str]]]:
    """Read PDF → (text with [IMAGE_N] placeholders, [(bytes, ext), ...])."""
    doc = fitz.open(stream=data, filetype="pdf")
    images: list[tuple[bytes, str]] = []
    img_counter = 0
    parts: list[str] = []

    for page in doc:
        page_text = page.get_text()
        # Extract images from this page
        page_images = page.get_images(full=True)
        if page_images:
            img_refs: list[str] = []
            for img_info in page_images:
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                except Exception:
                    continue
                if base_image and base_image.get("image"):
                    img_counter += 1
                    ext = base_image.get("ext", "png")
                    images.append((base_image["image"], ext))
                    img_refs.append(f"[IMAGE_{img_counter}]")
            if img_refs:
                page_text += "\n" + "\n".join(img_refs) + "\n"
        parts.append(page_text)

    doc.close()
    return "\n".join(parts), images


def read_file(filepath: Path) -> str:
    """Read file as plain text (no image extraction)."""
    suffix = filepath.suffix.lower()
    if suffix in (".md", ".txt", ".csv", ".log", ".conf", ".ini",
                   ".yaml", ".yml", ".json", ".xml"):
        return filepath.read_text(encoding="utf-8")
    elif suffix == ".docx":
        if not HAS_DOCX:
            print("❌ pip install python-docx", file=sys.stderr)
            sys.exit(1)
        doc = Document(str(filepath))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    elif suffix == ".pdf":
        if not HAS_PDF:
            print("❌ pip install pymupdf", file=sys.stderr)
            sys.exit(1)
        doc = fitz.open(str(filepath))
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    else:
        try:
            return filepath.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            print(f"❌ Format non supporté : {suffix}", file=sys.stderr)
            sys.exit(1)


def read_file_with_images(
    filepath: Path,
) -> tuple[str, list[tuple[bytes, str]]]:
    """Read file with image extraction.
    Returns (text_with_IMAGE_placeholders, [(image_bytes, extension), ...]).
    For text-only formats, images list is empty.
    """
    suffix = filepath.suffix.lower()
    if suffix == ".docx":
        if not HAS_DOCX:
            print("❌ pip install python-docx", file=sys.stderr)
            sys.exit(1)
        return _read_docx_with_images(filepath.read_bytes())
    elif suffix == ".pdf":
        if not HAS_PDF:
            print("❌ pip install pymupdf", file=sys.stderr)
            sys.exit(1)
        return _read_pdf_with_images(filepath.read_bytes())
    else:
        return read_file(filepath), []


def read_file_bytes(data: bytes, filename: str) -> str:
    """Read file from bytes in memory (text only, no image extraction)."""
    suffix = Path(filename).suffix.lower()
    if suffix in (".md", ".txt", ".csv", ".log", ".conf", ".ini",
                   ".yaml", ".yml", ".json", ".xml"):
        return data.decode("utf-8")
    elif suffix == ".docx":
        if not HAS_DOCX:
            raise ImportError("python-docx requis : pip install python-docx")
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    elif suffix == ".pdf":
        if not HAS_PDF:
            raise ImportError("pymupdf requis : pip install pymupdf")
        doc = fitz.open(stream=data, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text
    else:
        return data.decode("utf-8")


def read_file_bytes_with_images(
    data: bytes, filename: str,
) -> tuple[str, list[tuple[bytes, str]]]:
    """Read file from bytes with image extraction.
    Returns (text_with_IMAGE_placeholders, [(image_bytes, extension), ...]).
    """
    suffix = Path(filename).suffix.lower()
    if suffix == ".docx":
        if not HAS_DOCX:
            raise ImportError("python-docx requis : pip install python-docx")
        return _read_docx_with_images(data)
    elif suffix == ".pdf":
        if not HAS_PDF:
            raise ImportError("pymupdf requis : pip install pymupdf")
        return _read_pdf_with_images(data)
    else:
        return read_file_bytes(data, filename), []


def save_images(
    images: list[tuple[bytes, str]], output_dir: Path,
) -> list[str]:
    """Save extracted images to a directory.
    Images are named IMAGE_1.ext, IMAGE_2.ext, etc.
    Returns list of saved filenames.
    """
    if not images:
        return []
    output_dir.mkdir(parents=True, exist_ok=True)
    filenames: list[str] = []
    for i, (img_data, ext) in enumerate(images, 1):
        fname = f"IMAGE_{i}.{ext}"
        (output_dir / fname).write_bytes(img_data)
        filenames.append(fname)
    return filenames


# =============================================================================
# DÉCOUPAGE
# =============================================================================

def split_into_chunks(text: str, max_chars: int = 4000) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    paragraphs = re.split(r'\n\s*\n', text)
    chunks, current, current_len = [], [], 0
    for para in paragraphs:
        para_len = len(para) + 2
        if current_len + para_len > max_chars and current:
            chunks.append("\n\n".join(current))
            current, current_len = [para], para_len
        else:
            current.append(para)
            current_len += para_len
    if current:
        chunks.append("\n\n".join(current))
    final_chunks = []
    for chunk in chunks:
        if len(chunk) <= max_chars:
            final_chunks.append(chunk)
        else:
            sub_current, sub_len = [], 0
            for line in chunk.split("\n"):
                ll = len(line) + 1
                if sub_len + ll > max_chars and sub_current:
                    final_chunks.append("\n".join(sub_current))
                    sub_current, sub_len = [line], ll
                else:
                    sub_current.append(line)
                    sub_len += ll
            if sub_current:
                final_chunks.append("\n".join(sub_current))
    return final_chunks


# =============================================================================
# VÉRIFICATION POST-ANONYMISATION
# =============================================================================

def post_check(text: str) -> list[str]:
    warnings = []
    ips = re.findall(r'\b(?:\d{1,3}\.){3}\d{1,3}\b', text)
    if ips:
        warnings.append(
            f"IPs potentiellement restantes : {', '.join(set(ips))}"
        )
    emails = re.findall(
        r'\b[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}\b', text
    )
    if emails:
        warnings.append(
            f"Emails potentiellement restants : {', '.join(set(emails))}"
        )
    fqdns = re.findall(
        r'\b[a-zA-Z][a-zA-Z0-9\-]*\.'
        r'(?:[a-zA-Z0-9\-]+\.)*'
        r'(?:local|lan|internal|corp|intra)\b',
        text,
    )
    if fqdns:
        warnings.append(
            f"FQDN internes restants : {', '.join(set(fqdns))}"
        )
    return warnings


def count_llm_tags(text: str) -> dict[str, int]:
    full_tags = re.findall(
        r'\[(?:PERSONNE|ENTREPRISE|SITE|PROJET|LIEU|REF|SECRET)_\d+\]', text
    )
    tag_counts = defaultdict(int)
    for t in full_tags:
        tag_counts[t] += 1
    return dict(tag_counts)


# =============================================================================
# UTILITAIRES
# =============================================================================

def check_ollama(
    base_url: str = "http://localhost:11434",
    model: str = "gpt-oss:20b",
) -> tuple[bool, str, list[str]]:
    if not HAS_REQUESTS:
        return False, "Module 'requests' non installé", []
    try:
        r = requests.get(f"{base_url}/api/tags", timeout=5)
        r.raise_for_status()
        models = [m["name"] for m in r.json().get("models", [])]
        if any(model in m for m in models):
            return True, f"Ollama connecté — {model} prêt", models
        else:
            return True, f"Ollama connecté — {model} NON trouvé", models
    except requests.exceptions.ConnectionError:
        return False, "Ollama non joignable — lance 'ollama serve'", []
    except Exception as e:
        return False, f"Erreur Ollama : {e}", []


# =============================================================================
# PIPELINE PRINCIPAL
# =============================================================================

def _run_llm_pass(chunks, system_prompt, pass_name, log, model, ollama_url,
                  timeout, cancel_flag, on_progress,
                  done_chunks_ref, total_chunks):
    """Execute one LLM pass on a list of chunks. Returns result list."""
    log.stats["llm_passes"] += 1
    result_chunks = []
    for i, chunk in enumerate(chunks, 1):
        if cancel_flag and cancel_flag.is_set():
            log.log("STOP", "Annulé par l'utilisateur.")
            log.stats["annule"] = True
            result_chunks.append(chunk)
            continue
        progress = 0.2 + (done_chunks_ref[0] / total_chunks) * 0.65
        log.log(
            "LLM",
            f"  {pass_name} [{i}/{len(chunks)}] ({len(chunk)} chars)...",
            progress=progress,
        )
        result, success = call_ollama_chat(
            chunk, system_prompt,
            model=model, base_url=ollama_url, timeout=timeout,
        )
        if success:
            if result.strip() == chunk.strip():
                log.log(
                    "WARN",
                    f"  {pass_name} [{i}/{len(chunks)}] "
                    "LLM n'a fait aucun changement.",
                )
                log.stats["llm_no_change"] = (
                    log.stats.get("llm_no_change", 0) + 1
                )
            else:
                log.log(
                    "OK", f"  {pass_name} [{i}/{len(chunks)}] traité."
                )
        else:
            log.log(
                "ERROR" if "Passe 2" in pass_name else "WARN",
                f"  {pass_name} [{i}/{len(chunks)}] erreur — texte conservé.",
            )
            log.stats["llm_erreurs"] += 1
        result_chunks.append(result)
        log.stats["llm_chunks_traites"] += 1
        done_chunks_ref[0] += 1
        progress = 0.2 + (done_chunks_ref[0] / total_chunks) * 0.65
        if on_progress:
            on_progress(
                f"{pass_name} — chunk {i}/{len(chunks)} — {log.elapsed()}",
                progress,
            )
    return result_chunks


def run_pipeline(
    text: str,
    filename: str = "document",
    custom_words: dict[str, str] | None = None,
    use_llm: bool = True,
    model: str = "gpt-oss:20b",
    ollama_url: str = "http://localhost:11434",
    chunk_size: int = 4000,
    passes: int = 2,
    timeout: int = 300,
    on_progress: Callable[[str, float], None] | None = None,
    cancel_flag: threading.Event | None = None,
    images_count: int = 0,
    images_folder: str = "",
    deep_analysis: bool = False,
) -> dict:
    """
    Execute the full anonymization pipeline.
    Returns {text, mapping, report, warnings, stats}.
    """
    log = Logger(on_progress=on_progress)
    log.stats["fichier_source"] = filename
    log.stats["taille_originale"] = len(text)
    log.stats["images_trouvees"] = images_count

    log.log(
        "INFO",
        f"Début anonymisation de {filename} ({len(text)} caractères)",
        progress=0.0,
    )

    if images_count > 0:
        log.log("IMG", f"{images_count} image(s) extraite(s) du document.")

    # ── Passe 0 : Mots personnalisés ─────────────────────────
    regex_anon = RegexAnonymizer()

    if custom_words:
        log.log(
            "CUSTOM",
            f"Passe 0 : {len(custom_words)} mot(s) personnalisé(s)...",
            progress=0.05,
        )
        text, custom_count = apply_custom_words(text, custom_words, regex_anon)
        log.stats["custom_remplacements"] = custom_count
        log.log("CUSTOM", f"{custom_count} remplacement(s).")

    # ── Passe 1 : Regex ──────────────────────────────────────
    log.log("REGEX", "Passe 1 : Anonymisation regex...", progress=0.1)
    text = regex_anon.anonymize(text)
    log.stats["regex_remplacements"] = (
        len(regex_anon.mapping) - (len(custom_words) if custom_words else 0)
    )
    log.log(
        "REGEX", f"{log.stats['regex_remplacements']} patterns regex remplacés."
    )

    # ── Passes LLM ───────────────────────────────────────────
    llm_entity_map = {}
    cancelled = cancel_flag and cancel_flag.is_set()

    if use_llm and not cancelled:
        log.log(
            "LLM",
            f"Test connexion Ollama ({ollama_url})...",
            progress=0.15,
        )
        connected, msg, _ = check_ollama(ollama_url, model)
        if not connected:
            log.log("ERROR", f"Ollama : {msg}")
            log.log("WARN", "Fallback regex uniquement.")
            use_llm = False

    if use_llm and not cancelled:
        # Préfixe Reasoning: low pour le mode rapide (par défaut)
        reasoning_prefix = "" if deep_analysis else "Reasoning: low\n\n"
        prompt_p2 = reasoning_prefix + SYSTEM_PROMPT_PASS2
        prompt_p3 = reasoning_prefix + SYSTEM_PROMPT_PASS3

        mode_label = "approfondie" if deep_analysis else "rapide"
        log.log("OK", f"Ollama OK — {model} (analyse {mode_label}).",
                progress=0.18)
        chunks = split_into_chunks(text, chunk_size)
        total_chunks = len(chunks) * min(passes, 3)
        done_ref = [0]  # mutable for pass-by-reference

        # Passe 2
        result_chunks = _run_llm_pass(
            chunks, prompt_p2, "Passe 2", log, model, ollama_url,
            timeout, cancel_flag, on_progress, done_ref, total_chunks,
        )
        text = "\n\n".join(result_chunks)

        # Passe 3
        if passes >= 2 and not (cancel_flag and cancel_flag.is_set()):
            chunks2 = split_into_chunks(text, chunk_size)
            result_chunks2 = _run_llm_pass(
                chunks2, prompt_p3, "Passe 3 vérif", log, model,
                ollama_url, timeout, cancel_flag, on_progress,
                done_ref, total_chunks,
            )
            text = "\n\n".join(result_chunks2)

        # Passe 4
        if passes >= 3 and not (cancel_flag and cancel_flag.is_set()):
            chunks3 = split_into_chunks(text, chunk_size)
            result_chunks3 = _run_llm_pass(
                chunks3, prompt_p3, "Passe 4 strict", log, model,
                ollama_url, timeout, cancel_flag, on_progress,
                done_ref, total_chunks,
            )
            text = "\n\n".join(result_chunks3)

        llm_entity_map = count_llm_tags(text)
    elif not use_llm:
        log.log("INFO", "Passe LLM ignorée.")

    # ── Vérification finale ──────────────────────────────────
    log.log("VERIF", "Vérification post-anonymisation...", progress=0.9)
    warnings = post_check(text)
    log.stats["verif_warnings"] = warnings
    if warnings:
        for w in warnings:
            log.log("WARN", w)
    else:
        log.log("OK", "Aucun pattern résiduel détecté.")

    log.stats["taille_finale"] = len(text)
    log.stats["duree_totale"] = round(time.time() - log.start_time, 2)

    # ── Mapping ──────────────────────────────────────────────
    mapping_export = {}
    for key, tag in regex_anon.mapping.items():
        _, original = key.split("::", 1)
        mapping_export[tag] = original
    for tag in llm_entity_map:
        if tag not in mapping_export:
            mapping_export[tag] = "⟨détecté par LLM⟩"

    report = log.generate_report(
        regex_anon, llm_entity_map, images_folder=images_folder,
    )
    log.log("DONE", f"Terminé en {log.elapsed()}.", progress=1.0)

    return {
        "text": text,
        "mapping": mapping_export,
        "report": report,
        "warnings": warnings,
        "stats": log.stats,
        "regex_anon": regex_anon,
    }


# =============================================================================
# CLI
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Anonymisation hybride (Regex + LLM Ollama, multi-passe)",
        epilog="""
Exemples :
  python anonymize.py cahier_des_charges.docx
  python anonymize.py rapport.pdf --model gpt-oss:120b
  python anonymize.py notes.md -o anonyme.md --no-llm
  python anonymize.py spec.docx --passes 3 --chunk-size 2000""",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    parser.add_argument("fichier", help="Fichier à anonymiser")
    parser.add_argument("--model", default="gpt-oss:20b", help="Modèle Ollama")
    parser.add_argument("--output", "-o", help="Fichier de sortie")
    parser.add_argument("--ollama-url", default="http://localhost:11434")
    parser.add_argument("--no-llm", action="store_true",
                        help="Regex uniquement")
    parser.add_argument("--chunk-size", type=int, default=4000)
    parser.add_argument("--passes", type=int, default=2, choices=[1, 2, 3])
    parser.add_argument("--timeout", type=int, default=300)
    parser.add_argument(
        "--dict", metavar="FILE",
        help="Dictionnaire de mots sensibles (JSON). "
             "Défaut : sensitive-words.json à côté du script",
    )
    parser.add_argument(
        "--deep", action="store_true",
        help="Analyse approfondie — le LLM réfléchit plus "
             "(plus précis, plus lent)",
    )

    args = parser.parse_args()
    filepath = Path(args.fichier)

    if not filepath.exists():
        print(f"❌ Fichier introuvable : {filepath}", file=sys.stderr)
        sys.exit(1)

    print(f"\n{'='*60}")
    print(f"  ANONYMISATION — {filepath.name}")
    print(f"{'='*60}\n")

    # Read file + extract images
    text, images = read_file_with_images(filepath)

    images_folder_name = ""
    if images:
        images_dir = filepath.with_name(f"{filepath.stem}_images")
        saved = save_images(images, images_dir)
        images_folder_name = images_dir.name
        print(f"  🖼️  {len(images)} image(s) extraite(s) → {images_dir}")

    # Load dictionary (--dict or default sensitive-words.json)
    dict_path = Path(args.dict) if args.dict else None
    dict_words = load_sensitive_words(dict_path)
    if dict_words:
        print(f"  📖 Dictionnaire : {len(dict_words)} mot(s) chargé(s)")

    result = run_pipeline(
        text=text, filename=filepath.name,
        custom_words=dict_words if dict_words else None,
        use_llm=not args.no_llm, model=args.model,
        ollama_url=args.ollama_url, chunk_size=args.chunk_size,
        passes=args.passes, timeout=args.timeout,
        images_count=len(images),
        images_folder=images_folder_name,
        deep_analysis=args.deep,
    )

    output_path = (
        Path(args.output) if args.output
        else filepath.with_name(f"{filepath.stem}_anonymise.md")
    )
    output_path.write_text(result["text"], encoding="utf-8")

    mapping_path = filepath.with_name(f"{filepath.stem}_mapping.json")
    mapping_path.write_text(
        json.dumps(result["mapping"], indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    report_path = filepath.with_name(f"{filepath.stem}_rapport.md")
    report_path.write_text(result["report"], encoding="utf-8")

    s = result["stats"]
    print(f"\n{'='*60}")
    print(f"  RÉSUMÉ")
    print(f"{'='*60}")
    print(f"  📄 Source       : {filepath.name}")
    print(f"  📏 Taille       : {s['taille_originale']} → "
          f"{s['taille_finale']} chars")
    if s["custom_remplacements"]:
        print(f"  🏷️  Custom       : {s['custom_remplacements']}")
    print(f"  🔍 Regex        : {s['regex_remplacements']}")
    print(f"  🤖 LLM          : {s['llm_passes']} passes, "
          f"{s['llm_chunks_traites']} chunks")
    if images:
        print(f"  🖼️  Images       : {len(images)} → {images_folder_name}/")
    print(f"  ⏱️  Durée        : {s['duree_totale']}s")
    print(f"  ✅ Sortie       : {output_path}")
    print(f"{'='*60}\n")


if __name__ == "__main__":
    main()
